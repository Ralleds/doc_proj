import re
import docx
from datetime import datetime, timedelta
from docx.enum.text import WD_COLOR_INDEX


def increment_date(date_str):
    date_format = '%d.%m.%Y'  # Формат даты в вашем документе
    date = datetime.strptime(date_str, date_format)
    incremented_date = date + timedelta(days=1)
    return incremented_date.strftime(date_format)


def update_date_time(doc_path, new_date_time):
    doc = docx.Document(doc_path)

    # Изменяем дату и время в соответствии с заданным форматом
    for paragraph in doc.paragraphs:
        if "Дата и время посещения:" in paragraph.text:
            old_date_time = paragraph.text.split(": ")[1].strip()
            new_paragraph = paragraph.text.replace(old_date_time, new_date_time)
            paragraph.text = new_paragraph

            new_date_time = increment_date(new_date_time)

            # Установка размера шрифта 10 после изменения текста
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(10)

    doc.save(doc_path)


def update_name(doc_path, new_name):
    doc = docx.Document(doc_path)

    # Изменяем дату и время в соответствии с заданным форматом
    for paragraph in doc.paragraphs:
        if "Пациент:" in paragraph.text:
            old_name = paragraph.text.split(": ")[1].strip()
            new_paragraph = paragraph.text.replace(old_name, new_name)
            paragraph.text = new_paragraph

            # Установка размера шрифта 10 после изменения текста
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(10)

    doc.save(doc_path)


def update_date_birth(doc_path, new_birth):
    doc = docx.Document(doc_path)

    # Изменяем дату и время в соответствии с заданным форматом
    for paragraph in doc.paragraphs:
        if "Дата рождения:" in paragraph.text:
            old_birth = paragraph.text.split(": ")[1].strip()
            new_paragraph = paragraph.text.replace(old_birth, new_birth)
            paragraph.text = new_paragraph

            # Установка размера шрифта 10 после изменения текста
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(10)

    doc.save(doc_path)


def update_pos_date(doc_path, new_pos_date):
    doc = docx.Document(doc_path)

    # Изменяем дату и время в соответствии с заданным форматом
    for paragraph in doc.paragraphs:
        if "Актив на дому" in paragraph.text:
            old_date_time = paragraph.text.split("дому ")[1].strip()
            new_paragraph = paragraph.text.replace(old_date_time, new_pos_date)
            paragraph.text = new_paragraph

            new_pos_date = increment_date(new_pos_date)

            # Установка размера шрифта 10 после изменения текста
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(10)

    doc.save(doc_path)


def update_IMT(doc_path, data):
    doc = docx.Document(doc_path)

    def replace_height(matchobj):
        return f"Рост {data[0]}см"

    def replace_weight(matchobj):
        return f"вес {data[1]} кг"

    def replace_bmi(matchobj):
        bmi = int(data[1]) / ((int(data[0]) / 100) ** 2)
        formatted_bmi = "{:.1f}".format(bmi)

        return f"ИМТ {formatted_bmi.replace('.', ',')} кг"

    for paragraph in doc.paragraphs:
        if "Объективно" in paragraph.text:
            old_data = paragraph.text
            updated_str = re.sub(r"(Рост )(\d+)см", replace_height, old_data)
            updated_str = re.sub(r"(вес )(\d+) кг", replace_weight, updated_str)
            updated_str = re.sub(r"(ИМТ )(\d+,\d+) кг", replace_bmi, updated_str)
            paragraph.text = '' # empty string
            paragraph.add_run('Объективно: ').bold = True
            paragraph.add_run(updated_str[12:])

            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(10)

    doc.save(doc_path)


doc_path = "ex.docx"
new_name, new_birth, new_date_time, height, weight = input().split(';')
data = [height, weight]
print(new_name, new_birth, new_date_time, data, sep='\n')
new_pos_date = increment_date(new_date_time)

update_date_time(doc_path, new_date_time)
update_name(doc_path, new_name)
update_date_birth(doc_path, new_birth)
update_pos_date(doc_path, new_pos_date)
update_IMT(doc_path, data)
